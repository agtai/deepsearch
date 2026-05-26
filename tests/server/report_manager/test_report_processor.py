import re
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytest

from server.deepsearch.common.exception.exceptions import ReportConvertDependencyException
from server.deepsearch.core.manager.report_manager.conversion_utils import (
    ensure_pandoc,
    normalize_docx_tables,
    postprocess_html,
    preprocess_markdown_text,
    wrap_html_tables,
)
from server.deepsearch.core.manager.report_manager.docx_offline import convert_md_to_docx
from server.deepsearch.core.manager.report_manager.html_offline import convert_md_to_html
from server.deepsearch.core.manager.report_manager.mermaid_offline import (
    ensure_mermaid_cli,
    render_mermaid_offline,
)
from server.deepsearch.core.manager.report_manager.mermaid_preprocess import (
    MermaidRenderOptions,
    extract_xychart_metadata,
    preprocess_mermaid_code,
)
from server.deepsearch.core.manager.report_manager.report_processor import ReportHtml, ReportWord


def test_ensure_mermaid_cli_returns_unavailable_when_missing(monkeypatch):
    """Validate Mermaid CLI detection when the executable is unavailable.

    Args:
        monkeypatch: pytest monkeypatch fixture.

    Returns:
        None.
    """
    monkeypatch.delenv("MERMAID_MMDC_PATH", raising=False)
    monkeypatch.setattr("shutil.which", lambda name: None)
    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.mermaid_offline.resolve_mmdc_path",
        lambda: None,
    )

    status = ensure_mermaid_cli()

    assert status.available is False


def test_render_mermaid_offline_returns_false_when_cli_missing(tmp_path, monkeypatch):
    """Validate Mermaid rendering fallback when Mermaid CLI is missing.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    monkeypatch.delenv("MERMAID_MMDC_PATH", raising=False)
    monkeypatch.setattr("shutil.which", lambda name: None)
    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.mermaid_offline.resolve_mmdc_path",
        lambda: None,
    )

    ok = render_mermaid_offline(
        "graph TD\nA-->B",
        tmp_path / "diagram.svg",
        output_format="svg",
    )

    assert ok is False


def test_ensure_pandoc_raises_dependency_exception_when_download_fails(monkeypatch):
    """Validate pandoc setup failures are surfaced as dependency exceptions.

    Args:
        monkeypatch: pytest monkeypatch fixture.

    Returns:
        None.
    """
    monkeypatch.setattr("pypandoc.get_pandoc_version", lambda: (_ for _ in ()).throw(OSError("missing pandoc")))
    monkeypatch.setattr("pypandoc.download_pandoc", lambda: (_ for _ in ()).throw(RuntimeError("download failed")))

    with pytest.raises(ReportConvertDependencyException):
        ensure_pandoc()


def test_preprocess_mermaid_code_scales_xychart_and_extracts_metadata():
    """Validate xychart preprocessing keeps parity with the reference offline flow.

    Returns:
        None.
    """
    processed, supplement = preprocess_mermaid_code(
        "xychart-beta\n  bar [1200]\n",
        MermaidRenderOptions(),
    )
    metadata = extract_xychart_metadata(processed, warn_on_invalid=False)

    assert supplement == ""
    assert 'y-axis "x1e3"' in processed
    assert "bar [1.2]" in processed
    assert len(metadata.series) == 1
    assert metadata.series[0].display_values == ["1.2"]


def test_preprocess_markdown_text_strips_internal_citation_markers():
    """Validate export preprocessing hides internal citation control markers.

    Returns:
        None.
    """
    text = (
        "保留引用[checked_citation:4][[5]](https://example.com/source)\n\n"
        "移除旧标记[citation:2]"
    )

    processed = preprocess_markdown_text(text)

    assert "checked_citation" not in processed
    assert "[citation:2]" not in processed
    assert '[5]</a>' in processed


def test_wrap_html_tables_adds_centering_container_once():
    """Validate HTML table wrapping is idempotent.

    Returns:
        None.
    """
    html_text = "<p>intro</p><table><tr><td>A</td></tr></table>"

    processed = wrap_html_tables(html_text)
    processed_twice = wrap_html_tables(processed)

    assert '<div class="table-wrap"><table>' in processed
    assert processed.count('class="table-wrap"') == 1
    assert processed_twice.count('class="table-wrap"') == 1


def test_postprocess_html_wraps_tables_without_rewriting_svg():
    """Validate table wrapping does not mutate Mermaid SVG HTML.

    Returns:
        None.
    """
    html_text = (
        '<div class="mermaid-rendered"><svg viewBox="0 0 100 100"></svg></div>'
        "<table><tr><td>A</td></tr></table>"
    )

    processed = postprocess_html(html_text)

    assert 'viewBox="0 0 100 100"' in processed
    assert '<div class="table-wrap"><table>' in processed


def test_report_html_convert_from_markdown_wraps_tables():
    """Validate direct HTML conversion wraps Markdown tables.

    Returns:
        None.
    """
    html_text = ReportHtml.convert_from_markdown("| A | B |\n|---|---|\n| 1 | 2 |")

    assert 'class="table-wrap"' in html_text
    assert "<table>" in html_text


def test_report_word_convert_from_markdown_keeps_wrapped_tables():
    """Validate online DOCX conversion keeps tables wrapped for HTML centering.

    Returns:
        None.
    """
    doc = ReportWord.convert_from_markdown("| A | B |\n|---|---|\n| 1 | 2 |")

    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "A"
    assert doc.tables[0].cell(1, 1).text == "2"


def test_report_word_convert_from_html_handles_irregular_table_rows():
    """Validate HTML table conversion tolerates rows with different cell counts."""
    html_text = (
        '<div class="report-container">'
        "<table>"
        "<tr><th>A</th><th>B</th></tr>"
        "<tr><td>1</td><td>2</td><td>3</td></tr>"
        "</table>"
        "</div>"
    )

    doc = ReportWord._html_to_word(html_text)

    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 3
    assert doc.tables[0].cell(0, 2).text == ""
    assert doc.tables[0].cell(1, 2).text == "3"


def test_report_word_convert_from_html_limits_nested_block_depth():
    """Validate deeply nested HTML is flattened instead of recursing indefinitely."""
    html_text = (
        '<div class="report-container">'
        + "<div>" * 120
        + "<p>深层内容</p>"
        + "</div>" * 120
        + "</div>"
    )

    doc = ReportWord._html_to_word(html_text)

    assert any("深层内容" in paragraph.text for paragraph in doc.paragraphs)


def test_report_table_css_preserves_global_width_and_centers_wrapped_tables():
    """Validate report CSS limits table centering changes to wrapped tables.

    Returns:
        None.
    """
    css_path = Path("server/deepsearch/core/manager/report_manager/css/style.css")
    css_text = css_path.read_text(encoding="utf-8")

    assert re.search(r"table\s*\{[^}]*width:\s*100%;", css_text, flags=re.DOTALL)
    assert re.search(
        r"\.table-wrap\s+table\s*\{[^}]*width:\s*auto;[^}]*max-width:\s*100%;[^}]*margin:\s*0\s+auto;",
        css_text,
        flags=re.DOTALL,
    )


def test_report_html_export_renders_mermaid_or_falls_back(tmp_path):
    """Validate HTML export renders Mermaid or preserves source as fallback.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    final_result = {
        "response_content": "# 标题\n\n```mermaid\ngraph TD\nA-->B\n```",
        "infer_messages": [],
        "chart_messages": [],
        "warning_info": "",
        "exception_info": "",
    }

    html_text = ReportHtml.convert_from_final_result(final_result, tmp_path)

    assert "<html" in html_text.lower()
    assert "标题" in html_text
    assert ("<svg" in html_text) or ("language-mermaid" in html_text)


def test_normalize_docx_tables_centers_tables_and_captions(tmp_path):
    """Validate DOCX table normalization centers table objects and captions.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    docx_path = tmp_path / "tables.docx"
    document = Document()
    document.add_paragraph("普通正文段落")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.add_paragraph("表2-1：合肥市“三电”系统核心企业的技术实力与市场表现")
    document.save(docx_path)

    normalize_docx_tables(docx_path)

    normalized = Document(docx_path)
    assert normalized.tables[0].alignment == WD_TABLE_ALIGNMENT.CENTER
    assert normalized.paragraphs[0].alignment is None
    assert normalized.paragraphs[-1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_convert_md_to_html_annotates_xychart_value_labels(tmp_path, monkeypatch):
    """Validate HTML export annotates xychart SVG output with value labels.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "```mermaid\nxychart-beta\n  bar [1200]\n```",
        encoding="utf-8",
    )

    def _fake_render_mermaid_offline(code, output_path, **kwargs):
        del code, kwargs
        output_file = tmp_path / output_path.name
        output_file.write_text(
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100">'
            '<g class="plot"><g class="bar-plot-0" fill="#374151">'
            '<rect x="10" y="10" width="20" height="30" />'
            "</g></g></svg>",
            encoding="utf-8",
        )
        return True

    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.html_offline.render_mermaid_offline",
        _fake_render_mermaid_offline,
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "xychart-value-label" in html_text


def test_convert_md_to_html_keeps_legacy_font_caption_separate_from_following_list(tmp_path):
    """Validate legacy font captions do not absorb following bullet lists."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "![图表示例](chart.png)\n"
        "<font size=2>**图表示例**: 图注说明</font>[[1]](https://example.com)\n"
        "- **技术维度**：第一条\n"
        "- **经济维度**：第二条\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert ".figure-caption {" in html_text
    assert "width: 100%;" in html_text
    assert "text-align: center;" in html_text
    assert '<div class="figure-caption">' in html_text
    assert "<ul>" in html_text
    assert "<li><strong>技术维度</strong>" in html_text


def test_convert_md_to_html_separates_paragraph_from_following_bullets_without_blank_line(tmp_path):
    """Validate list items render after a paragraph even when the source misses a blank line."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "在代际价值观层面，这一人群展现出强烈的矛盾统一体特征：\n"
        "- **求稳与求变并存**：第一条\n"
        "- **务实与悦己交织**：第二条\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "<p>在代际价值观层面，这一人群展现出强烈的矛盾统一体特征：</p>" in html_text
    assert "<ul>" in html_text
    assert "<li><strong>求稳与求变并存</strong>：第一条</li>" in html_text


def test_convert_md_to_html_centers_table_display(tmp_path):
    """Validate exported HTML uses centered table presentation styles."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "| 列1 | 列2 |\n"
        "| --- | --- |\n"
        "| A | B |\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "table {" in html_text
    assert "margin: 16px auto 24px;" in html_text
    assert "width: fit-content;" in html_text
    assert "max-width: 100%;" in html_text
    assert "th[style], td[style] {" in html_text
    assert "text-align: center !important;" in html_text
    assert "text-align: center;" in html_text


def test_report_docx_export_creates_docx_file(tmp_path, monkeypatch):
    """Validate DOCX export writes a pandoc-generated file into the bundle.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    final_result = {
        "response_content": "# 标题\n\n```mermaid\ngraph TD\nA-->B\n```",
        "infer_messages": [],
        "chart_messages": [],
        "warning_info": "",
        "exception_info": "",
    }

    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.ensure_pandoc",
        lambda: None,
    )
    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.normalize_docx_fonts",
        lambda *_args, **_kwargs: None,
    )
    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.normalize_docx_tables",
        lambda *_args, **_kwargs: None,
    )

    def _fake_convert_file(*args, **kwargs):
        outputfile = kwargs["outputfile"]
        with open(outputfile, "wb") as file:
            file.write(b"PK\x03\x04docx")

    monkeypatch.setattr("pypandoc.convert_file", _fake_convert_file)

    docx_path = ReportWord.convert_from_final_result(final_result, tmp_path)

    assert docx_path.exists()
    assert docx_path.read_bytes().startswith(b"PK")


def test_convert_md_to_docx_normalizes_headings_fonts_and_tables(tmp_path, monkeypatch):
    """Validate DOCX export uses heading/font/table post-processing flow.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text("1. 一级标题\n", encoding="utf-8")

    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.ensure_pandoc",
        lambda: None,
    )

    captured: dict[str, str] = {}

    def _fake_convert_file(input_file, *_args, **kwargs):
        captured["content"] = Path(input_file).read_text(encoding="utf-8")
        Path(kwargs["outputfile"]).write_bytes(b"PK\x03\x04docx")

    font_calls = {"count": 0}
    table_calls = {"count": 0}

    monkeypatch.setattr("pypandoc.convert_file", _fake_convert_file)
    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.normalize_docx_fonts",
        lambda *_args, **_kwargs: font_calls.__setitem__("count", font_calls["count"] + 1),
        raising=False,
    )
    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.normalize_docx_tables",
        lambda *_args, **_kwargs: table_calls.__setitem__("count", table_calls["count"] + 1),
        raising=False,
    )

    convert_md_to_docx(md_path, docx_path)

    assert '<h1 id="1">1 一级标题</h1>' in captured["content"]
    assert font_calls["count"] == 1
    assert table_calls["count"] == 1


def test_convert_md_to_docx_preserves_short_bold_spans_in_long_chinese_summary(tmp_path):
    """Validate DOCX export keeps inline bold spans in long Chinese summary paragraphs."""
    docx_module = pytest.importorskip("docx")
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        pytest.skip("pandoc is unavailable in the current test environment")

    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# 摘要\n\n"
        "都市年轻职场人群（22-35岁）在职业周期演进与高居住成本制约下呈现显著分化，"
        "其核心生存图景由**69.6%**企业新招毕业生硕士占比更高、一线城市高达**45.6%**"
        "的房租负担率及**79.2%**企业离职率低于**10%**的求稳心态共同刻画；该群体深陷"
        "工作高压（**31.5%**日工作超10小时）、时间剥夺（北京单程通勤**47分钟**）、社交"
        "萎缩（超三成频繁孤独）与生活品质坍塌（超**90%**受亚健康影响）交织的痛点因果网，"
        "驱动其行为向情绪补剂常态化（近九成为情绪买单）、社交模块化（**54.4%**有搭子）"
        "与技术双刃剑（超**56%**日常使用GenAI但超六成担忧失业）三大策略演化；由此催生"
        "效率工具（**56.1%**愿为AI付费）、情绪价值（四成向AI倾诉）、零糖社交与零家务闭环"
        "等复合需求，其消费决策呈现极致折叠的精算师特质（比价工具使用率达**78%**）与为情绪"
        "溢价买单并存，复购核心由体验确证（**77.6%**因使用感好复购）与情绪持续供给驱动。\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = docx_module.Document(docx_path)
    summary_paragraph = document.paragraphs[1]
    bold_runs = {run.text for run in summary_paragraph.runs if run.text and run.bold}

    assert "**" not in summary_paragraph.text
    assert {
        "69.6%",
        "45.6%",
        "79.2%",
        "10%",
        "31.5%",
        "47分钟",
        "90%",
        "54.4%",
        "56%",
        "56.1%",
        "78%",
        "77.6%",
    }.issubset(bold_runs)


def test_report_docx_export_raises_dependency_exception_when_pandoc_setup_fails(tmp_path, monkeypatch):
    """Validate DOCX export propagates pandoc dependency failures.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    final_result = {
        "response_content": "# 标题",
        "infer_messages": [],
        "chart_messages": [],
        "warning_info": "",
        "exception_info": "",
    }

    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.ensure_pandoc",
        lambda: (_ for _ in ()).throw(ReportConvertDependencyException("pandoc missing")),
    )

    with pytest.raises(ReportConvertDependencyException):
        ReportWord.convert_from_final_result(final_result, tmp_path)


def test_report_docx_export_raises_dependency_exception_when_pandoc_execution_fails(tmp_path, monkeypatch):
    """Validate DOCX export maps pandoc runtime dependency failures.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    final_result = {
        "response_content": "# 标题",
        "infer_messages": [],
        "chart_messages": [],
        "warning_info": "",
        "exception_info": "",
    }

    monkeypatch.setattr(
        "server.deepsearch.core.manager.report_manager.docx_offline.ensure_pandoc",
        lambda: None,
    )
    monkeypatch.setattr(
        "pypandoc.convert_file",
        lambda *args, **kwargs: (_ for _ in ()).throw(OSError("pandoc execution failed")),
    )

    with pytest.raises(ReportConvertDependencyException):
        ReportWord.convert_from_final_result(final_result, tmp_path)
